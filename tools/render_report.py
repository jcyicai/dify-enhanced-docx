import io
import json
import re
from collections.abc import Generator
from typing import Any

import markdown
import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
from bs4 import BeautifulSoup, NavigableString, Tag
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FONT_NAME = "Microsoft YaHei"


class RenderReportDocxTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        title = str(tool_parameters.get("cover_title") or "研究报告").strip()
        subtitle = str(tool_parameters.get("cover_subtitle") or "").strip()
        meta = str(tool_parameters.get("cover_meta") or "").strip()
        content = str(tool_parameters.get("markdown_content") or "").strip()
        filename = self._safe_filename(
            str(tool_parameters.get("output_filename") or title).strip()
        )

        document = Document()
        self._configure_page(document)
        self._configure_styles(document)
        self._add_cover(document, title, subtitle, meta)
        document.add_page_break()
        self._add_toc_page(document, content)
        document.add_page_break()
        self._add_markdown(document, content)
        self._enable_field_update(document)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        yield self.create_blob_message(
            blob=buffer.read(),
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{filename}.docx",
            },
        )

    def _configure_page(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    def _configure_styles(self, document: Document) -> None:
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = FONT_NAME
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_after = Pt(6)

        heading_settings = {
            "Title": (26, 0, 18),
            "Subtitle": (15, 0, 12),
            "Heading 1": (18, 20, 10),
            "Heading 2": (15, 16, 8),
            "Heading 3": (12, 12, 6),
            "Heading 4": (11, 10, 4),
        }
        for style_name, (size, before, after) in heading_settings.items():
            style = styles[style_name]
            style.font.name = FONT_NAME
            style.font.size = Pt(size)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        code = styles.add_style("Report Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(8)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.space_before = Pt(10)
        code.paragraph_format.space_after = Pt(14)

        toc = styles.add_style("Static TOC", WD_STYLE_TYPE.PARAGRAPH)
        toc.font.name = FONT_NAME
        toc.font.size = Pt(11)
        toc._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        toc.paragraph_format.line_spacing = 1.5
        toc.paragraph_format.space_after = Pt(5)

    def _add_cover(self, document: Document, title: str, subtitle: str, meta: str) -> None:
        for _ in range(5):
            document.add_paragraph()

        p = document.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        self._set_run_font(run, 26)

        if subtitle:
            p = document.add_paragraph(style="Subtitle")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            self._set_run_font(run, 15)

        for _ in range(7):
            document.add_paragraph()

        for line in meta.splitlines():
            if not line.strip():
                continue
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line.strip())
            self._set_run_font(run, 12)

    def _add_toc_page(self, document: Document, markdown_content: str) -> None:
        p = document.add_paragraph(style="Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("目录")
        run.bold = True
        self._set_run_font(run, 22)

        headings = self._extract_headings(markdown_content)
        for level, text in headings:
            p = document.add_paragraph(style="Static TOC")
            p.paragraph_format.left_indent = Cm(max(0, level - 1) * 0.7)
            run = p.add_run(text)
            self._set_run_font(run, 11 if level == 1 else 10.5)
            if level == 1:
                run.bold = True

    def _extract_headings(self, content: str) -> list[tuple[int, str]]:
        headings: list[tuple[int, str]] = []
        for line in content.splitlines():
            match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
            if match:
                text = re.sub(r"[*_`]", "", match.group(2)).strip()
                headings.append((len(match.group(1)), text))
        return headings

    def _add_markdown(self, document: Document, content: str) -> None:
        html = markdown.markdown(
            content,
            extensions=["tables", "fenced_code", "sane_lists"],
        )
        soup = BeautifulSoup(html, "html.parser")
        first_h1 = True
        for element in soup.children:
            if not isinstance(element, Tag):
                continue
            if element.name in {"h1", "h2", "h3", "h4"}:
                level = int(element.name[1])
                if level == 1 and not first_h1:
                    document.add_page_break()
                if level == 1:
                    first_h1 = False
                p = document.add_paragraph(style=f"Heading {level}")
                self._append_inline(p, element)
            elif element.name == "p":
                p = document.add_paragraph()
                self._append_inline(p, element)
            elif element.name in {"ul", "ol"}:
                self._add_list(document, element)
            elif element.name == "table":
                self._add_table(document, element)
            elif element.name == "pre":
                code_tag = element.find("code")
                classes = code_tag.get("class", []) if code_tag else []
                if code_tag and "language-echarts" in classes:
                    try:
                        option = json.loads(code_tag.get_text())
                        image = self._render_echarts(option)
                        p = document.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(10)
                        p.paragraph_format.space_after = Pt(16)
                        run = p.add_run()
                        run.add_picture(image, width=Cm(16.2))
                    except Exception:
                        p = document.add_paragraph(style="Report Code")
                        p.add_run(element.get_text())
                else:
                    p = document.add_paragraph(style="Report Code")
                    p.add_run(element.get_text())
            elif element.name == "blockquote":
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                self._append_inline(p, element)
            elif element.name == "hr":
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)

    def _append_inline(self, paragraph, element: Tag) -> None:
        for child in element.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(str(child))
                self._set_run_font(run, 10.5)
            elif isinstance(child, Tag):
                text = child.get_text()
                run = paragraph.add_run(text)
                self._set_run_font(run, 10.5)
                if child.name in {"strong", "b"}:
                    run.bold = True
                if child.name in {"em", "i"}:
                    run.italic = True
                if child.name == "code":
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)

    def _add_list(self, document: Document, element: Tag) -> None:
        style = "List Number" if element.name == "ol" else "List Bullet"
        for li in element.find_all("li", recursive=False):
            p = document.add_paragraph(style=style)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(4)
            self._append_inline(p, li)

    def _add_table(self, document: Document, element: Tag) -> None:
        rows = element.find_all("tr")
        if not rows:
            return
        col_count = max(len(row.find_all(["th", "td"])) for row in rows)
        table = document.add_table(rows=0, cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(rows):
            cells = table.add_row().cells
            for col_index, cell_tag in enumerate(row.find_all(["th", "td"])):
                cell = cells[col_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.text = cell_tag.get_text(strip=True)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                    for run in p.runs:
                        self._set_run_font(run, 9)
                        if row_index == 0:
                            run.bold = True
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)

    def _enable_field_update(self, document: Document) -> None:
        settings = document.settings._element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")

    def _render_echarts(self, option: dict[str, Any]) -> io.BytesIO:
        plt.rcParams["font.sans-serif"] = [
            "Microsoft YaHei",
            "Noto Sans CJK SC",
            "SimHei",
            "Arial Unicode MS",
            "DejaVu Sans",
        ]
        plt.rcParams["axes.unicode_minus"] = False

        grids = option.get("grid") or [{}]
        count = len(grids)
        if count == 4:
            rows, cols = 2, 2
            figsize = (13, 9)
        elif count == 2:
            rows, cols = 1, 2
            figsize = (13, 5.5)
        else:
            rows, cols = count, 1
            figsize = (12, max(4.5, 3.5 * count))

        fig, axes = plt.subplots(rows, cols, figsize=figsize, squeeze=False)
        flat_axes = list(axes.flat)
        titles = option.get("title") or []
        if isinstance(titles, dict):
            titles = [titles]
        overall_title = titles[0].get("text", "") if titles else ""
        subplot_titles = [item.get("text", "") for item in titles[1:]]

        x_axes = option.get("xAxis") or []
        y_axes = option.get("yAxis") or []
        if isinstance(x_axes, dict):
            x_axes = [x_axes]
        if isinstance(y_axes, dict):
            y_axes = [y_axes]

        y_by_grid: dict[int, list[int]] = {}
        for index, axis in enumerate(y_axes):
            y_by_grid.setdefault(int(axis.get("gridIndex", 0)), []).append(index)

        axis_cache: dict[tuple[int, int], Any] = {}
        for grid_index in range(count):
            ax = flat_axes[grid_index]
            if grid_index < len(subplot_titles):
                ax.set_title(subplot_titles[grid_index], fontsize=12, pad=12)
            y_indices = y_by_grid.get(grid_index, [grid_index])
            for order, y_index in enumerate(y_indices):
                axis_cache[(grid_index, y_index)] = ax if order == 0 else ax.twinx()

        for series in option.get("series") or []:
            x_index = int(series.get("xAxisIndex", 0))
            y_index = int(series.get("yAxisIndex", 0))
            grid_index = int(
                (x_axes[x_index] if x_index < len(x_axes) else {}).get("gridIndex", 0)
            )
            ax = axis_cache.get((grid_index, y_index), flat_axes[grid_index])
            x_axis = x_axes[x_index] if x_index < len(x_axes) else {}
            y_axis = y_axes[y_index] if y_index < len(y_axes) else {}
            self._plot_series(ax, series, x_axis, y_axis)

        for grid_index in range(count):
            ax = flat_axes[grid_index]
            ax.grid(True, alpha=0.22)
            ax.tick_params(axis="x", labelrotation=25)
            handles, labels = ax.get_legend_handles_labels()
            for y_index in y_by_grid.get(grid_index, [])[1:]:
                twin = axis_cache.get((grid_index, y_index))
                if twin is not None:
                    h2, l2 = twin.get_legend_handles_labels()
                    handles.extend(h2)
                    labels.extend(l2)
            if handles:
                ax.legend(handles, labels, fontsize=8, loc="best")

        if overall_title:
            fig.suptitle(overall_title, fontsize=16, fontweight="bold", y=0.99)
        fig.tight_layout(rect=(0, 0, 1, 0.96), h_pad=3.0, w_pad=2.5)
        output = io.BytesIO()
        fig.savefig(output, format="png", dpi=170, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        output.seek(0)
        return output

    def _plot_series(self, ax, series: dict[str, Any], x_axis: dict[str, Any], y_axis: dict[str, Any]) -> None:
        chart_type = series.get("type", "line")
        name = series.get("name", "")
        data = series.get("data") or []
        x_categories = x_axis.get("data") or []
        y_categories = y_axis.get("data") or []

        if chart_type == "heatmap":
            width = len(x_categories)
            height = len(y_categories)
            matrix = np.full((height, width), np.nan)
            for x, y, value in data:
                matrix[int(y), int(x)] = float(value)
            image = ax.imshow(matrix, aspect="auto", cmap="RdYlGn", origin="upper", vmin=np.nanmin(matrix), vmax=np.nanmax(matrix))
            ax.set_xticks(range(width), labels=x_categories)
            ax.set_yticks(range(height), labels=y_categories)
            for y in range(height):
                for x in range(width):
                    if not np.isnan(matrix[y, x]):
                        ax.text(x, y, f"{matrix[y, x]:.1f}", ha="center", va="center", fontsize=7)
            ax.figure.colorbar(image, ax=ax, fraction=0.035, pad=0.03)
            return

        if chart_type == "boxplot":
            stats = []
            for index, values in enumerate(data):
                if len(values) < 5:
                    continue
                stats.append(
                    {
                        "label": x_categories[index] if index < len(x_categories) else str(index + 1),
                        "whislo": values[0],
                        "q1": values[1],
                        "med": values[2],
                        "q3": values[3],
                        "whishi": values[4],
                        "fliers": [],
                    }
                )
            ax.bxp(stats, showfliers=False, patch_artist=True, boxprops={"facecolor": "#8FAEC5", "alpha": 0.8})
            return

        if data and isinstance(data[0], (list, tuple)):
            xs = [item[0] for item in data]
            ys = [item[1] for item in data]
        else:
            xs = list(range(len(data)))
            ys = data
            if x_categories:
                ax.set_xticks(xs, labels=x_categories)

        if chart_type == "scatter":
            ax.scatter(xs, ys, s=18, alpha=0.5, label=name)
        elif chart_type == "bar":
            bars = ax.bar(xs, ys, label=name, alpha=0.85)
            labels = (series.get("label") or {}).get("show")
            if labels:
                ax.bar_label(bars, fmt="%.1f", padding=3, fontsize=8)
        else:
            ax.plot(xs, ys, linewidth=1.8, label=name)

        if x_axis.get("name"):
            ax.set_xlabel(x_axis["name"])
        if y_axis.get("name"):
            ax.set_ylabel(y_axis["name"])

    def _set_run_font(self, run, size: float) -> None:
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    def _safe_filename(self, value: str) -> str:
        value = re.sub(r'[\\/:*?"<>|]', "_", value)
        return value[:120] or "report"
